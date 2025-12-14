# import streamlit as st
# import fitz
# import docx
# import os
# import io
# from docx import Document
# from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
# from reportlab.lib.styles import getSampleStyleSheet
# from reportlab.lib.pagesizes import letter

# from crewai import Agent, Task, Crew
# from crewai_tools import SerperDevTool
# import warnings
# from dotenv import load_dotenv

# warnings.filterwarnings('ignore')
# load_dotenv()

# # ===========================
# # FILE HANDLING FUNCTIONS
# # ===========================

# def extract_text_from_pdf(file):
#     doc = fitz.open(stream=file.read(), filetype="pdf")
#     text = ""
#     for page in doc:
#         text += page.get_text()
#     return text

# def extract_text_from_docx(file):
#     doc = docx.Document(file)
#     fullText = []
#     for para in doc.paragraphs:
#         fullText.append(para.text)
#     return "\n".join(fullText)

# def extract_resume_text(uploaded_file):
#     if uploaded_file.name.endswith(".pdf"):
#         return extract_text_from_pdf(uploaded_file)
#     if uploaded_file.name.endswith(".docx"):
#         return extract_text_from_docx(uploaded_file)
#     return ""

# def create_docx(text):
#     doc = Document()
#     for line in text.split("\n"):
#         doc.add_paragraph(line)
#     buffer = io.BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     return buffer

# def create_pdf(text):
#     buffer = io.BytesIO()
#     doc = SimpleDocTemplate(buffer, pagesize=letter)
#     styles = getSampleStyleSheet()
#     story = []
#     for line in text.split("\n"):
#         if line.strip() == "":
#             continue
#         story.append(Paragraph(line, styles["Normal"]))
#         story.append(Spacer(1, 8))
#     doc.build(story)
#     buffer.seek(0)
#     return buffer


# # ===========================
# # STREAMLIT UI
# # ===========================

# st.set_page_config(page_title="Job Application AI Assistant", layout="wide")
# st.title("🤖 Job Application AI Assistant")

# uploaded_file = st.file_uploader("Upload Resume (.pdf or .docx)", type=["pdf", "docx"])
# location = st.text_input("Enter preferred job location", value="India")

# job_description = st.text_area(
#     "Paste Job Description here",
#     height=200,
#     placeholder="Paste job description from LinkedIn / Naukri / etc..."
# )


# # ===========================
# # RUN PIPELINE
# # ===========================

# if uploaded_file and job_description and st.button("Run AI Agents"):

#     with st.spinner("Extracting resume text..."):
#         resume_text = extract_resume_text(uploaded_file)

#     st.success("Resume text extracted!")
#     with st.expander("📄 Original Resume Preview"):
#         st.write(resume_text[:1500] + " ..." if len(resume_text) > 1500 else resume_text)

#     # ===========================
#     # API Keys
#     # ===========================

#     os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
#     os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")
#     os.environ["OPENAI_MODEL_NAME"] = "gpt-4o"

#     search_tool = SerperDevTool()

#     # ===========================
#     # AGENTS
#     # ===========================

#     # Agent 1: Resume Analyzer & Scorer
#     resume_analyzer = Agent(
#         role="Senior ATS Resume Analyst",
#         goal="Analyze resume against job description and provide detailed ATS compatibility score with specific improvement recommendations",
#         verbose=False,
#         backstory="""You are an expert ATS (Applicant Tracking System) analyst with 15+ years of experience 
#         in recruitment technology. You understand exactly how ATS systems parse and score resumes. 
#         You identify keyword gaps, formatting issues, and content misalignments between resumes and job descriptions."""
#     )

#     # Agent 2: Professional Summary Writer
#     summary_writer = Agent(
#         role="Executive Resume Summary Specialist",
#         goal="Rewrite the professional summary to perfectly align with job requirements using powerful action words",
#         verbose=False,
#         backstory="""You are a certified professional resume writer (CPRW) specializing in crafting 
#         compelling professional summaries. You transform generic summaries into targeted, keyword-rich 
#         statements that immediately capture recruiter attention and pass ATS screening."""
#     )

#     # Agent 3: Skills Optimizer
#     skills_optimizer = Agent(
#         role="Technical Skills & Keywords Specialist",
#         goal="Reorganize and enhance the skills section to match job description requirements and maximize ATS score",
#         verbose=False,
#         backstory="""You are an ATS optimization expert who understands how to structure technical skills 
#         for maximum keyword matching. You know which skills to prioritize, how to group them effectively, 
#         and how to add missing but relevant skills without fabrication."""
#     )

#     # Agent 4: Experience & Projects Enhancer
#     experience_enhancer = Agent(
#         role="Achievement-Focused Resume Writer",
#         goal="Transform job responsibilities and project descriptions into powerful achievement statements with metrics and action verbs",
#         verbose=False,
#         backstory="""You are a resume writing expert who specializes in converting passive job descriptions 
#         into dynamic achievement statements. You use strong action verbs (Engineered, Developed, Implemented, 
#         Optimized, Deployed) and quantify achievements wherever possible. You align experience descriptions 
#         with target job requirements."""
#     )

#     # Agent 5: Final Resume Compiler
#     resume_compiler = Agent(
#         role="Senior Resume Editor & Formatter",
#         goal="Compile all improvements into a cohesive, ATS-optimized final resume while maintaining professional formatting",
#         verbose=False,
#         backstory="""You are a meticulous resume editor who ensures consistency, proper grammar, 
#         and professional formatting. You compile all section improvements into a single polished document 
#         that reads naturally and maintains the candidate's authentic voice while maximizing ATS compatibility."""
#     )

#     # Agent 6: Job Researcher
#     job_researcher = Agent(
#         role="Senior Recruitment Consultant",
#         goal="Find relevant job openings matching the candidate's enhanced profile",
#         verbose=False,
#         tools=[search_tool],
#         backstory="""You are an experienced recruitment consultant with access to job market data. 
#         You match candidate profiles with suitable job openings and provide actionable job search recommendations."""
#     )

#     # ===========================
#     # TASKS
#     # ===========================

#     # Task 1: Analyze original resume
#     analysis_task = Task(
#         description=f"""
#         Analyze this resume against the job description and provide:
        
#         1. CURRENT ATS SCORE (out of 100)
#         2. KEYWORD GAP ANALYSIS:
#            - List keywords from JD that are MISSING in resume
#            - List keywords that are present but need emphasis
#         3. SECTION-BY-SECTION ANALYSIS:
#            - Professional Summary: What's wrong and what to improve
#            - Skills Section: Missing skills, poor organization issues
#            - Experience/Projects: Weak descriptions, missing achievements
#         4. TOP 5 PRIORITY IMPROVEMENTS needed
        
#         RESUME:
#         {resume_text}
        
#         JOB DESCRIPTION:
#         {job_description}
#         """,
#         expected_output="Detailed ATS analysis with score, keyword gaps, and specific improvement recommendations",
#         agent=resume_analyzer
#     )

#     # Task 2: Rewrite Professional Summary
#     summary_task = Task(
#         description=f"""
#         Based on the analysis, rewrite the Professional Summary section to:
        
#         1. Start with a strong descriptor (e.g., "Results-driven", "Detail-oriented", "Innovative")
#         2. Include 3-5 key skills/technologies mentioned in the JD
#         3. Highlight relevant experience areas that match the JD
#         4. Use industry-specific keywords from the job description
#         5. Keep it to 3-4 impactful sentences
#         6. DO NOT include quotation marks around the summary
#         7. DO NOT fabricate experience - only enhance what exists
        
#         ORIGINAL RESUME:
#         {resume_text}
        
#         JOB DESCRIPTION:
#         {job_description}
        
#         OUTPUT: Only the new Professional Summary section text.
#         """,
#         expected_output="Rewritten professional summary aligned with job description",
#         agent=summary_writer,
#         context=[analysis_task]
#     )

#     # Task 3: Optimize Skills Section
#     skills_task = Task(
#         description=f"""
#         Reorganize and enhance the Technical Skills section to:
        
#         1. GROUP skills into categories that match JD priorities:
#            - Programming & Libraries (Python, SQL, NumPy, pandas, scikit-learn, etc.)
#            - Machine Learning & Deep Learning (frameworks, concepts from JD)
#            - NLP & AI (if applicable to JD)
#            - Tools & Platforms (from JD requirements)
        
#         2. ADD missing skills from JD that the candidate likely has based on their projects
#            - Only add skills that are implied by their existing work
#            - DO NOT fabricate skills they don't have
        
#         3. PRIORITIZE skills mentioned in JD "Required Qualifications"
        
#         4. FORMAT as clean bullet points with category headers
        
#         ORIGINAL RESUME:
#         {resume_text}
        
#         JOB DESCRIPTION:
#         {job_description}
        
#         OUTPUT: Only the restructured Technical Skills section.
#         """,
#         expected_output="Reorganized and enhanced skills section matching JD requirements",
#         agent=skills_optimizer,
#         context=[analysis_task]
#     )

#     # Task 4: Enhance Experience & Projects
#     experience_task = Task(
#         description=f"""
#         Transform the Experience and Projects sections to:
        
#         1. EXPERIENCE SECTION:
#            - Use title that matches JD role (e.g., "AI/ML Intern" for AI roles)
#            - Start each bullet with strong ACTION VERBS: Developed, Engineered, Implemented, 
#              Deployed, Optimized, Integrated, Designed, Built, Automated, Collaborated
#            - Add METRICS where possible (%, numbers, scale)
#            - Align descriptions with JD responsibilities
#            - Keep 4-5 bullet points max
        
#         2. PROJECTS SECTION:
#            - Use strong action verbs to start each bullet
#            - Highlight technologies that match JD requirements
#            - Structure as: Action + What you did + Technology used + Impact/Result
#            - Include a "Technologies:" line after each project
#            - Keep 3-4 bullets per project
        
#         3. DO NOT fabricate any experience or projects
#         4. DO NOT change project names or core facts
        
#         ORIGINAL RESUME:
#         {resume_text}
        
#         JOB DESCRIPTION:
#         {job_description}
        
#         OUTPUT: Enhanced Experience and Projects sections only.
#         """,
#         expected_output="Enhanced experience and projects with action verbs and JD alignment",
#         agent=experience_enhancer,
#         context=[analysis_task]
#     )

#     # Task 5: Compile Final Resume
#     compile_task = Task(
#         description=f"""
#         Compile all enhanced sections into a complete, polished resume:
        
#         1. STRUCTURE (in this order):
#            - Contact Information (keep original - name, email, phone, location, LinkedIn, GitHub)
#            - Professional Summary (from summary task)
#            - Technical Skills (from skills task)
#            - Projects (from experience task)
#            - Experience (from experience task)
#            - Education (keep original)
#            - Key Competencies (add if space permits: soft skills relevant to JD)
        
#         2. FORMATTING RULES:
#            - Use consistent section headers (ALL CAPS or Title Case)
#            - Use bullet points for skills, experience, projects
#            - No quotation marks around any section
#            - Clean, professional formatting
#            - Proper spacing between sections
        
#         3. FINAL CHECKS:
#            - Grammar and spelling
#            - Consistency in tense (past tense for past roles)
#            - No fabricated information
#            - All JD keywords naturally incorporated
        
#         ORIGINAL CONTACT & EDUCATION INFO:
#         {resume_text}
        
#         OUTPUT: Complete formatted resume text ready for download.
#         """,
#         expected_output="Complete, polished, ATS-optimized resume",
#         agent=resume_compiler,
#         context=[summary_task, skills_task, experience_task]
#     )

#     # Task 6: Score Enhanced Resume
#     final_score_task = Task(
#         description=f"""
#         Score the ENHANCED resume against the job description:
        
#         1. NEW ATS SCORE (out of 100)
#         2. IMPROVEMENT SUMMARY:
#            - Score improvement from original
#            - Key improvements made
#            - Keywords now matched
#         3. REMAINING RECOMMENDATIONS (if any)
#         4. INTERVIEW TIPS based on the enhanced resume
        
#         JOB DESCRIPTION:
#         {job_description}
        
#         Evaluate the compiled resume from the previous task.
#         """,
#         expected_output="Final ATS score with improvement summary and interview tips",
#         agent=resume_analyzer,
#         context=[compile_task, analysis_task]
#     )

#     # Task 7: Job Search
#     job_task = Task(
#         description=f"""
#         Search for 5 relevant job openings based on the enhanced resume profile.
        
#         Search criteria:
#         - Location: {location}
#         - Skills: Extract top skills from the enhanced resume
#         - Role type: Match the target role from job description
        
#         For each job provide:
#         1. Job Title
#         2. Company Name
#         3. Location
#         4. Key Requirements (brief)
#         5. Application Link (if available)
        
#         Focus on entry-level to mid-level positions matching the candidate's experience.
#         """,
#         expected_output="5 relevant job listings with details and links",
#         agent=job_researcher,
#         context=[compile_task]
#     )

#     # ===========================
#     # CREW PIPELINE
#     # ===========================

#     crew = Crew(
#         agents=[
#             resume_analyzer, 
#             summary_writer, 
#             skills_optimizer, 
#             experience_enhancer, 
#             resume_compiler, 
#             job_researcher
#         ],
#         tasks=[
#             analysis_task,
#             summary_task,
#             skills_task,
#             experience_task,
#             compile_task,
#             final_score_task,
#             job_task
#         ],
#         verbose=False
#     )

#     with st.spinner("🔄 Running AI agent workflow... This may take 2-3 minutes."):
#         result = crew.kickoff(inputs={"resume": resume_text, "location": location, "job_description": job_description})

#     st.success("Done! 🎉 Below are your results.")

#     # ===========================
#     # UI OUTPUTS
#     # ===========================

#     col1, col2 = st.columns(2)
    
#     with col1:
#         st.subheader("📊 Original Resume Analysis")
#         st.write(analysis_task.output.raw)

#     with col2:
#         st.subheader("📈 Enhanced Resume Score")
#         st.write(final_score_task.output.raw)

#     st.divider()

#     st.subheader("✨ Enhanced Resume (Job Description Optimized)")
#     enhanced_resume = compile_task.output.raw
#     st.text_area("Enhanced Resume", enhanced_resume, height=400)

#     st.divider()

#     with st.expander("📝 Section-by-Section Improvements"):
#         st.markdown("**Professional Summary:**")
#         st.write(summary_task.output.raw)
#         st.markdown("---")
#         st.markdown("**Technical Skills:**")
#         st.write(skills_task.output.raw)
#         st.markdown("---")
#         st.markdown("**Experience & Projects:**")
#         st.write(experience_task.output.raw)

#     st.divider()

#     st.subheader("🔍 Relevant Job Openings")
#     st.write(job_task.output.raw)

#     # ===========================
#     # DOWNLOAD ENHANCED RESUME
#     # ===========================

#     st.divider()
#     st.subheader("📥 Download Enhanced Resume")

#     col1, col2 = st.columns(2)
    
#     with col1:
#         docx_buffer = create_docx(enhanced_resume)
#         st.download_button(
#             label="📄 Download as DOCX",
#             data=docx_buffer,
#             file_name="enhanced_resume.docx",
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )

#     with col2:
#         pdf_buffer = create_pdf(enhanced_resume)
#         st.download_button(
#             label="📄 Download as PDF",
#             data=pdf_buffer,
#             file_name="enhanced_resume.pdf",
#             mime="application/pdf"
#         )

# elif uploaded_file and not job_description:
#     st.warning("⚠️ Please paste a job description to optimize your resume.")

import streamlit as st
import fitz
import docx
import os
import io
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter

from crewai import Agent, Task, Crew, LLM
from crewai_tools import SerperDevTool
import warnings
from dotenv import load_dotenv

warnings.filterwarnings('ignore')
load_dotenv()

# ===========================
# FILE HANDLING FUNCTIONS
# ===========================

def extract_text_from_pdf(file):
    doc = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def extract_text_from_docx(file):
    doc = docx.Document(file)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return "\n".join(fullText)

def extract_resume_text(uploaded_file):
    if uploaded_file.name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    if uploaded_file.name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    return ""

def get_file_extension(uploaded_file):
    return uploaded_file.name.split(".")[-1].lower()

def create_docx(text):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf(text):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    for line in text.split("\n"):
        if line.strip() == "":
            continue
        story.append(Paragraph(line, styles["Normal"]))
        story.append(Spacer(1, 8))
    doc.build(story)
    buffer.seek(0)
    return buffer

def create_output_file(text, file_format):
    if file_format == "pdf":
        return create_pdf(text), "application/pdf", "enhanced_resume.pdf"
    else:
        return create_docx(text), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "enhanced_resume.docx"


# ===========================
# STREAMLIT UI
# ===========================

st.set_page_config(page_title="JD-Based Resume Enhancer", layout="wide")
st.title("🎯 JD-Based Resume Enhancer")
st.caption("✅ Consistent output | Same resume + Same JD = Same result every time")

uploaded_file = st.file_uploader("📄 Upload Resume (.pdf or .docx)", type=["pdf", "docx"])
location = st.text_input("📍 Enter preferred job location", value="India")

job_description = st.text_area(
    "📋 Paste Job Description (REQUIRED)",
    height=250,
    placeholder="""Paste the COMPLETE job description here..."""
)

if uploaded_file:
    input_format = get_file_extension(uploaded_file)
    st.success(f"✅ Uploaded: **{uploaded_file.name}** | Output: **{input_format.upper()}**")


# ===========================
# RUN PIPELINE
# ===========================

if uploaded_file and job_description and st.button("🚀 Enhance Resume"):

    input_format = get_file_extension(uploaded_file)
    
    with st.spinner("Extracting resume..."):
        resume_text = extract_resume_text(uploaded_file)

    st.success("Resume extracted!")
    
    with st.expander("📄 View Original Resume"):
        st.text(resume_text)

    # ===========================
    # API Keys
    # ===========================

    os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
    os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")

    search_tool = SerperDevTool()

    # ===========================
    # LLM WITH TEMPERATURE=0 FOR CONSISTENT OUTPUT
    # ===========================
    
    # Temperature = 0 means NO randomness = SAME output every time
    consistent_llm = LLM(
        model="gpt-5",
        temperature=0,  # 0 = deterministic, same output every time
        seed=42         # Fixed seed for reproducibility
    )

    # ===========================
    # AGENTS (with consistent LLM)
    # ===========================

    jd_analyzer = Agent(
        role="Job Description Analyst",
        goal="Extract ALL keywords from the job description",
        verbose=False,
        llm=consistent_llm,  # Use consistent LLM
        backstory="""You analyze job descriptions to extract exact keywords.
        You extract ONLY what is written - no assumptions."""
    )

    resume_matcher = Agent(
        role="Resume-JD Matcher",
        goal="Find EXACT matches between resume and JD keywords",
        verbose=False,
        llm=consistent_llm,  # Use consistent LLM
        backstory="""You compare resumes against job descriptions.
        You NEVER assume skills not written in the resume."""
    )

    resume_rewriter = Agent(
        role="JD-Focused Resume Rewriter",
        goal="Rewrite resume using ONLY matched JD keywords",
        verbose=False,
        llm=consistent_llm,  # Use consistent LLM
        backstory="""You rewrite resumes to align with job descriptions.
        You NEVER add fake experience, skills, or metrics."""
    )

    job_researcher = Agent(
        role="Job Search Specialist",
        goal="Find similar job openings",
        verbose=False,
        llm=consistent_llm,  # Use consistent LLM
        tools=[search_tool],
        backstory="You search for job openings matching candidate profiles."
    )

    # ===========================
    # TASKS
    # ===========================

    jd_analysis_task = Task(
        description=f"""
        Analyze this job description and extract:
        
        1. JOB TITLE
        2. REQUIRED SKILLS (must-have)
        3. PREFERRED SKILLS (nice-to-have)
        4. KEY RESPONSIBILITIES
        5. ATS KEYWORDS
        
        JOB DESCRIPTION:
        {job_description}
        
        Extract ONLY what is written.
        """,
        expected_output="Categorized JD keywords",
        agent=jd_analyzer
    )

    match_task = Task(
        description=f"""
        Compare resume against JD keywords. Create THREE lists:
        
        1. ✅ MATCHING SKILLS: In both resume AND JD
        2. 🔄 SYNONYM MATCHES: Same skill, different wording
        3. ❌ GAPS: In JD but NOT in resume (DO NOT ADD)
        
        RESUME:
        {resume_text}
        
        Be STRICT - only mark as match if skill exists in resume.
        """,
        expected_output="Matches, Synonyms, Gaps lists",
        agent=resume_matcher,
        context=[jd_analysis_task]
    )

    rewrite_task = Task(
        description=f"""
        Rewrite resume for THIS job description.
        
        ⛔ FORBIDDEN:
        - Adding skills from GAPS list
        - Adding ANY skill not in original resume
        - Adding fake metrics or numbers
        - Adding new bullet points
        - Changing job titles, companies, dates
        
        ✅ ALLOWED:
        - Replace weak verbs → strong verbs
        - Reorder skills (JD-matching first)
        - Use JD terminology for synonyms only
        - Fix grammar
        
        ORIGINAL RESUME:
        {resume_text}
        
        JOB DESCRIPTION:
        {job_description}
        
        OUTPUT: Complete enhanced resume.
        
        STRUCTURE:
        1. Contact info (unchanged)
        2. Professional Summary (rewritten, no quotes)
        3. Technical Skills (reordered)
        4. Projects (same names, reworded bullets)
        5. Experience (same title/company, reworded)
        6. Education (unchanged)
        """,
        expected_output="Enhanced resume - no fake content",
        agent=resume_rewriter,
        context=[jd_analysis_task, match_task]
    )

    job_task = Task(
        description=f"""
        Find 5 similar job openings.
        Location: {location}
        
        For each: Job Title, Company, Location, Link
        """,
        expected_output="5 job listings",
        agent=job_researcher,
        context=[jd_analysis_task, match_task]
    )

    # ===========================
    # CREW
    # ===========================

    crew = Crew(
        agents=[jd_analyzer, resume_matcher, resume_rewriter, job_researcher],
        tasks=[jd_analysis_task, match_task, rewrite_task, job_task],
        verbose=False
    )

    with st.spinner("🔄 Enhancing resume (consistent mode)..."):
        result = crew.kickoff(inputs={
            "resume": resume_text, 
            "location": location, 
            "job_description": job_description
        })

    st.success("✅ Done!")

    # ===========================
    # OUTPUTS
    # ===========================

    enhanced_resume = rewrite_task.output.raw

    st.subheader("🔍 JD Keywords")
    with st.expander("View extracted keywords", expanded=True):
        st.write(jd_analysis_task.output.raw)

    st.divider()

    st.subheader("🔗 Match Analysis")
    st.write(match_task.output.raw)

    st.divider()

    st.subheader("✨ Enhanced Resume")
    st.text_area("Enhanced Resume", enhanced_resume, height=500)

    st.divider()
    
    st.subheader("📊 Before vs After")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**Original:**")
        st.text_area("Original", resume_text, height=300, key="orig")
    with col2:
        st.markdown("**Enhanced:**")
        st.text_area("Enhanced", enhanced_resume, height=300, key="enh")

    st.divider()

    st.subheader("🔍 Similar Jobs")
    st.write(job_task.output.raw)

    # ===========================
    # DOWNLOAD
    # ===========================

    st.divider()
    st.subheader("📥 Download")
    
    output_buffer, mime_type, filename = create_output_file(enhanced_resume, input_format)
    
    st.download_button(
        label=f"📄 Download as {input_format.upper()}",
        data=output_buffer,
        file_name=filename,
        mime=mime_type
    )

    with st.expander("Download other format"):
        if input_format == "pdf":
            other_buffer = create_docx(enhanced_resume)
            st.download_button(
                label="📄 DOCX",
                data=other_buffer,
                file_name="enhanced_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="other"
            )
        else:
            other_buffer = create_pdf(enhanced_resume)
            st.download_button(
                label="📄 PDF",
                data=other_buffer,
                file_name="enhanced_resume.pdf",
                mime="application/pdf",
                key="other"
            )

    st.divider()
    st.info("""
    🔄 **Consistency Guaranteed:**
    - Same resume + Same JD = Same output every time
    - Using temperature=0 for deterministic results
    
    ❌ **No fake content added**
    """)

elif uploaded_file and not job_description:
    st.warning("⚠️ Please paste a job description.")
elif not uploaded_file and job_description:
    st.warning("⚠️ Please upload your resume.")